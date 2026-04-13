"""
Document text extraction for TXT, PDF, and DOCX files.
All parsers return plain UTF-8 text.
"""

import io
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Routes to the correct parser based on file extension."""

    def parse(self, file_bytes: bytes, extension: str) -> str:
        """
        Extract text from file bytes.
        extension: '.txt', '.pdf', or '.docx'
        Returns: plain text string
        Raises: ValueError on parse failure
        """
        ext = extension.lower()
        if ext == '.txt':
            return self._parse_txt(file_bytes)
        elif ext == '.pdf':
            return self._parse_pdf(file_bytes)
        elif ext == '.docx':
            return self._parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported extension: {ext}")

    @staticmethod
    def _parse_txt(file_bytes: bytes) -> str:
        """Attempt multiple encodings for TXT files."""
        for encoding in ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1', 'ascii']:
            try:
                text = file_bytes.decode(encoding)
                logger.debug(f"TXT decoded with {encoding}")
                return text
            except (UnicodeDecodeError, ValueError):
                continue
        raise ValueError("Could not decode text file. Ensure it is saved as UTF-8.")

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for page_num, page in enumerate(doc):
                text = page.get_text("text").strip()
                if text:
                    pages.append(text)
                else:
                    logger.debug(f"Page {page_num + 1} has no extractable text (may be image).")
            doc.close()
        except Exception as e:
            raise ValueError(f"Failed to open PDF: {e}")

        if not pages:
            raise ValueError(
                "No extractable text found in PDF. "
                "The file may contain only scanned images. "
                "Please use a text-based PDF."
            )

        return "\n\n".join(pages)

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        """Extract paragraph text from DOCX."""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text and text not in paragraphs:
                            paragraphs.append(text)

        except Exception as e:
            raise ValueError(f"Failed to open DOCX: {e}")

        if not paragraphs:
            raise ValueError("No extractable text found in DOCX.")

        return "\n\n".join(paragraphs)
