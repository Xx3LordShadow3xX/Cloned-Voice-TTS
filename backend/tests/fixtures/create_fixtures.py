"""
Script to create test fixture files.
Run once: python tests/fixtures/create_fixtures.py
"""
import os

fixtures_dir = os.path.join(os.path.dirname(__file__))

# sample.txt
with open(os.path.join(fixtures_dir, 'sample.txt'), 'w', encoding='utf-8') as f:
    f.write("This is a sample text file for testing.\n\n")
    f.write("It contains multiple paragraphs.\n\n")
    f.write("Unicode test: café, résumé, naïve, über, 日本語\n\n")
    f.write("The quick brown fox jumps over the lazy dog.\n")

# sample.pdf
try:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Sample PDF content for testing.")
    page.insert_text((50, 80), "Page 1 of the test document.")
    doc.save(os.path.join(fixtures_dir, 'sample.pdf'))
    doc.close()
    print("Created sample.pdf")
except ImportError:
    print("Skipping sample.pdf (PyMuPDF not available)")

# sample.docx
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Sample DOCX content for testing.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(os.path.join(fixtures_dir, 'sample.docx'))
    print("Created sample.docx")
except ImportError:
    print("Skipping sample.docx (python-docx not available)")

# empty.txt
with open(os.path.join(fixtures_dir, 'empty.txt'), 'wb') as f:
    pass  # Zero bytes

print("Fixtures created in:", fixtures_dir)
