"""Unit tests for DocumentParser."""
import pytest
from services.parser import DocumentParser


parser = DocumentParser()


class TestTXTParsing:
    def test_utf8(self):
        text = parser.parse("Hello, world!".encode('utf-8'), '.txt')
        assert "Hello, world!" in text

    def test_latin1(self):
        text = parser.parse("Café résumé".encode('latin-1'), '.txt')
        assert text  # Should decode without error

    def test_empty_returns_empty_string(self):
        # Empty bytes decode successfully to "" — the parser returns "".
        # The upstream synthesize endpoint raises HTTP 422 on empty text.
        result = parser.parse(b"", '.txt')
        assert result == ""

    def test_utf8_bom(self):
        text = parser.parse("\ufeffHello".encode('utf-8-sig'), '.txt')
        assert "Hello" in text


class TestPDFParsing:
    def test_valid_pdf(self, tmp_path):
        # Minimal PDF with text (use fixture file)
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "Test PDF content")
        pdf_bytes = doc.tobytes()
        doc.close()

        text = parser.parse(pdf_bytes, '.pdf')
        assert "Test PDF content" in text

    def test_empty_pdf_raises(self):
        import fitz
        doc = fitz.open()
        doc.new_page()  # Page with no text
        pdf_bytes = doc.tobytes()
        doc.close()

        with pytest.raises(ValueError, match="No extractable text"):
            parser.parse(pdf_bytes, '.pdf')

    def test_corrupt_pdf_raises(self):
        with pytest.raises(ValueError):
            parser.parse(b"not a pdf at all", '.pdf')


class TestDOCXParsing:
    def test_valid_docx(self):
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("Test DOCX content")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text = parser.parse(docx_bytes, '.docx')
        assert "Test DOCX content" in text

    def test_empty_docx_raises(self):
        from docx import Document
        import io
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        with pytest.raises(ValueError, match="No extractable text"):
            parser.parse(docx_bytes, '.docx')


class TestUnsupportedExtension:
    def test_unknown_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported extension"):
            parser.parse(b"data", '.exe')
